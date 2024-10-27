# Костыль
from docx import Document
from docxtpl import DocxTemplate
import csv
from collections import defaultdict


with open("marathon.csv", 'r', encoding="utf-8") as file:
    data = csv.reader(file, delimiter=',')
    if not data:
        print("Файл пустой!")
        exit()

    csv_data = list()
    for row in data:
        csv_data.append(row)

data = list(sorted(csv_data, key=lambda x: x[0], reverse=True))  # Сортирует года

year_groups = defaultdict(list)  # Создает тип данных

# Группирует данные по годам
for item in data:
    year = item[0]
    year_groups[year].append(item)

for year, year_data in sorted(year_groups.items(), reverse=True): # Сортирует данные по городам внутри каждого года
    year_data.sort(key=lambda x: x[-1])

csv_data = list(year_groups.values())  # Делает данные обратно в формат csv файла
csv_data = [item for t in csv_data for item in t]  # Список списков -> сплошной список

new_document = Document()
existed_year = {}

# Создание шаблона

# enumerate - перечисляет текущую итерацию
for num, i in enumerate(csv_data):  # num - номер строки
  
    if num != 0:  # На первой итерации нет предыдущего города
        previous_city = csv_data[num-1][-1]
    else:
        previous_city = None

    if key in existed_year:  # Такой год уже есть
        line = existed_year[key]
        if current_city == previous_city:  # Добавляет человека к городу существующему
            line.add_run("\n{{{{ {human} }}}} - {{{{ {name} }}}} - {{{{ {time} }}}}".format(human=human_type, name=name, time=time))
        else:  # Создает новый город
            line.add_run("\n\nГород: {{{{ {city} }}}}\n{{{{ {human} }}}} - {{{{ {name} }}}} - {{{{ {time} }}}}".format(city=city, human=human_type, name=name, time=time))

    else:
        line = new_document.add_paragraph("Год: {{{{ {year} }}}}\n\nГород: {{{{ {city} }}}}\n{{{{ {human} }}}} - {{{{ {name} }}}} - {{{{ {time} }}}}".format(year=year, city=city, human=human_type, name=name, time=time))

        existed_year[key] = line
        new_document.add_page_break()  # Новая страница

new_document.save("sample.docx")

doc = DocxTemplate("sample.docx")

context = dict()
for num, i in enumerate(csv_data):  # Заполняем данными шаблон
    year = f"year_{num}"
    city = f"city_{num}"
    name = f"name_{num}"
    time = f"time_{num}"
    human_type = f"human_{num}"
    key = f"{i[0]}"

    context[year] = i[0]
    context[city] = i[-1]
    context[time] = i[-2]
    context[name] = i[1]
    context[human_type] = i[2]

doc.render(context)
doc.save("marathon.docx")
