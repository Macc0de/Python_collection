from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

new_document = Document()

heading = new_document.add_heading("Студент Маринич Максим", level=0)  # Заголовок
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

new_line = new_document.add_paragraph().add_run()  # Пустая строка

# 1 параграф
line = new_document.add_paragraph("Находится в городе Ярославль. ")
line_format = line.runs[0]
line_format.bold = True
line_format.italic = True
line_format.font.name = "Arial"
line_format.font.color.rgb = RGBColor(220, 20, 60)
line_format.font.size = Pt(20)

line.paragraph_format.line_spacing = Mm(10)  # Межстрочный интервал в 1 параграфе

sub_line = line.add_run("Проживает в общежитии.")  # Часть первого параграфа
sub_line.bold = True
sub_line.underline = True
sub_line.font.name = "Times New Roman"
sub_line.font.size = Pt(15)
sub_line.font.color.rgb = RGBColor(210, 105, 30)

line.paragraph_format.space_after = Mm(10)  # Расстоянием между 1 и 2 параграфом

# 2 параграф
line2 = new_document.add_paragraph("Учится в ЯрГУ на программиста.")  # Новый параграф с новой строки
line_format2 = line2.runs[0]
line_format2.italic = True
line_format2.font.name = "Cascadia Code"
line_format2.font.size = Pt(13)
line_format2.font.color.rgb = RGBColor(0, 0, 128)

line2.paragraph_format.space_before = Mm(5)

# 3 параграф
line3 = new_document.add_paragraph("Летом ездит на море.")
line3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
line_format3 = line3.runs[0]
line_format3.underline = True
line_format3.font.name = "Trebuchet MS"
line_format3.font.size = Pt(16)
line_format3.font.color.rgb = RGBColor(0, 255, 0)

paragraph_image = new_document.add_paragraph()  # Центр картинки делается через параграф
run = paragraph_image.add_run()
image = run.add_picture('cat2.jpg', width=Pt(300))
paragraph_image.alignment = WD_ALIGN_PARAGRAPH.CENTER

new_document.save("student.docx")
